#!/usr/bin/env python3
"""Genera Word de gastos referencia desde respaldo JSON de Cerebrus (exportJSON)."""
import json
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
except ImportError:
    print("Instala: pip install python-docx")
    sys.exit(1)

SECTION_KEYS = ["mercaderia", "rrhh", "operacion", "pasivos", "socias", "gestion"]
SEC_LABELS = {
    "mercaderia": "Mercadería",
    "rrhh": "RRHH",
    "operacion": "Operación",
    "pasivos": "Pasivos",
    "socias": "Socias",
    "gestion": "Gestión",
}
WDAYS = ["—", "Lun", "Mar", "Mié", "Jue", "Vie", "Sáb", "Dom"]


def load_state(path: Path) -> dict:
    raw = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(raw, dict) and "data" in raw:
        data = raw["data"]
        return data.get("state", data) if isinstance(data, dict) else {}
    if isinstance(raw, dict) and "state" in raw:
        return raw["state"]
    return raw if isinstance(raw, dict) else {}


def fecha_pago_item(item: dict) -> str:
    freq = int(item.get("freq") or 1)
    day = int(item.get("day") or 0)
    wd = int(item.get("weekday") or 0)
    if freq == 1 and day > 0:
        return f"Día {day} de cada mes"
    if freq == 4 and 0 < wd < len(WDAYS):
        return f"Cada {WDAYS[wd]}"
    if freq == 2:
        return "Quincenal"
    if freq == 30:
        return "Diario"
    return ""


def fmt_cl(n) -> str:
    n = int(round(float(n or 0)))
    return f"${n:,}".replace(",", ".")


def collect_fijos(state: dict) -> list:
    rows = []
    for k in SECTION_KEYS:
        for item in state.get(k) or []:
            name = (item.get("name") or "").strip()
            if not name:
                continue
            rows.append({
                "nombre": name,
                "seccion": SEC_LABELS.get(k, k),
                "fecha": fecha_pago_item(item),
                "monto": item.get("amount") or 0,
            })
    for g in (state.get("sebasFinanzas") or {}).get("gastosFijos") or []:
        name = (g.get("item") or "").strip()
        if not name:
            continue
        dia = g.get("dia")
        rows.append({
            "nombre": name,
            "seccion": "Finanzas Sebas · " + (g.get("cat") or ""),
            "fecha": f"Día {dia} del mes" if dia else "",
            "monto": g.get("valor") or 0,
        })
    return sorted(rows, key=lambda r: (r["seccion"], r["nombre"]))


def collect_otros(state: dict) -> list:
    otros = {}
    sac = state.get("sacCargas") or {}
    for period_key, arr in sac.items():
        if not isinstance(arr, list):
            continue
        for c in arr:
            if not c or c.get("_esPago"):
                continue
            name = (c.get("nombre") or "").strip()
            if not name:
                continue
            key = name.lower()
            m = max(c.get("monto") or 0, c.get("montoReal") or 0, c.get("pagadoParcial") or 0)
            if key not in otros:
                otros[key] = {"nombre": name, "cat": c.get("cat") or "", "monto": m, "periodos": []}
            otros[key]["monto"] = max(otros[key]["monto"], m)
            if period_key not in otros[key]["periodos"]:
                otros[key]["periodos"].append(period_key)
    return sorted(otros.values(), key=lambda r: r["nombre"])


def add_table(doc, headers, rows, empty_col_idx=4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val) if val is not None else ""
            if i == empty_col_idx:
                cells[i].text = " "
    return table


def build_doc(state: dict, out: Path):
    doc = Document()
    doc.add_heading("Gastos de referencia — CEREBRUS · Aquí Jaime", 0)
    doc.add_paragraph(
        f"Generado: {datetime.now().strftime('%d-%m-%Y %H:%M')} · "
        "Complete la columna «Monto nuevo»."
    )

    doc.add_heading("1. Gastos fijos (configuración y fechas de pago)", level=2)
    fijos = collect_fijos(state)
    fijos_rows = [
        [r["nombre"], r["seccion"], r["fecha"] or "—", fmt_cl(r["monto"]) if r["monto"] else "—", ""]
        for r in fijos
    ]
    if not fijos_rows:
        doc.add_paragraph("Sin ítems en configuración.")
    else:
        add_table(doc, ["Gasto", "Sección", "Fecha de pago", "Monto referencia", "Monto nuevo"], fijos_rows)

    doc.add_heading("2. Otros gastos (cargas en semanas guardadas)", level=2)
    doc.add_paragraph("Sin duplicar nombres. Monto referencia = mayor monto visto en SAC.")
    otros = collect_otros(state)
    otros_rows = [
        [
            r["nombre"],
            r["cat"] or "—",
            "; ".join(r["periodos"][:4]) + ("…" if len(r["periodos"]) > 4 else ""),
            fmt_cl(r["monto"]) + " (máx.)" if r["monto"] else "—",
            "",
        ]
        for r in otros
    ]
    if not otros_rows:
        doc.add_paragraph("Sin cargas en semanas guardadas.")
    else:
        add_table(doc, ["Gasto", "Categoría", "Períodos", "Monto referencia", "Monto nuevo"], otros_rows)

    doc.save(out)
    print(f"OK: {out} ({len(fijos)} fijos, {len(otros)} otros)")


def main():
    if len(sys.argv) < 2:
        print("Uso: python3 generar_gastos_referencia_word.py <Cerebrus_Respaldo.json> [salida.docx]")
        sys.exit(1)
    src = Path(sys.argv[1])
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_name("Cerebrus_Gastos_referencia.docx")
    state = load_state(src)
    build_doc(state, out)


if __name__ == "__main__":
    main()
